#!/usr/bin/env python3
"""
LaTeX to Word CV Converter
Converts LaTeX CV to Word document with high fidelity formatting preservation.
"""

import subprocess
import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from PIL import Image
import re

class LaTeXToWordConverter:
    def __init__(self, tex_file="cv.tex", docx_file="cv.docx"):
        self.tex_file = tex_file
        self.docx_file = docx_file
        self.doc = None
        
    def convert_with_pandoc(self):
        """Initial conversion using Pandoc"""
        print("Converting LaTeX to DOCX using Pandoc...")
        try:
            # Get absolute paths
            tex_abs_path = os.path.abspath(self.tex_file)
            temp_docx_path = os.path.join(os.path.dirname(self.docx_file), "temp_" + os.path.basename(self.docx_file))
            
            print(f"📁 Tex file: {tex_abs_path}")
            print(f"📁 Temp output: {temp_docx_path}")
            
            # Use Pandoc to convert LaTeX to DOCX
            cmd = [
                "pandoc",
                tex_abs_path,
                "-o", temp_docx_path,
                "--from", "latex",
                "--to", "docx",
                "--standalone"
            ]
            
            print(f"🔄 Running command: {' '.join(cmd)}")
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode != 0:
                print(f"❌ Pandoc error: {result.stderr}")
                print(f"📝 Pandoc stdout: {result.stdout}")
                return False
                
            print("✓ Pandoc conversion completed")
            return True
            
        except FileNotFoundError:
            print("Error: Pandoc not found. Please install Pandoc first.")
            return False
    
    def load_and_refine_document(self):
        """Load the Pandoc-generated document and refine formatting"""
        print("Loading and refining document formatting...")
        
        # Load the document
        temp_docx_path = os.path.join(os.path.dirname(self.docx_file), "temp_" + os.path.basename(self.docx_file))
        print(f"📁 Loading temp document: {temp_docx_path}")
        self.doc = Document(temp_docx_path)
        
        # Apply LaTeX-style formatting
        self._setup_document_styles()
        self._apply_page_layout()
        self._format_sections()
        self._format_lists()
        self._handle_header_with_photo()
        self._format_hyperlinks()
        self._fix_text_issues()
        
        print("✓ Document formatting refined")
    
    def _setup_document_styles(self):
        """Set up custom styles matching LaTeX formatting"""
        styles = self.doc.styles
        
        # Create or modify Normal style
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Liberation Serif'  # Closest to Computer Modern
        normal_font.size = Pt(10)
        
        # Set paragraph formatting
        normal_paragraph = normal_style.paragraph_format
        normal_paragraph.line_spacing = 0.9
        normal_paragraph.space_before = Pt(0)
        normal_paragraph.space_after = Pt(1)
        normal_paragraph.first_line_indent = Pt(0)
        
        # Create section style
        try:
            section_style = styles['Heading 1']
        except KeyError:
            section_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        
        section_font = section_style.font
        section_font.name = 'Liberation Serif'
        section_font.size = Pt(12)
        section_font.bold = True
        
        section_paragraph = section_style.paragraph_format
        section_paragraph.space_before = Pt(3)
        section_paragraph.space_after = Pt(1)
        section_paragraph.line_spacing = 1.0
    
    def _apply_page_layout(self):
        """Apply LaTeX page layout settings"""
        sections = self.doc.sections
        for section in sections:
            # Set margins to 1cm (matching LaTeX geometry)
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)
            
            # Set page size to A4
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
    
    def _format_sections(self):
        """Format section headers to match LaTeX titlesec"""
        for paragraph in self.doc.paragraphs:
            text = paragraph.text.strip()
            
            # Check if this is a section header (starts with common German CV sections)
            section_keywords = [
                "Kurzprofil", "Berufserfahrung", "Projekte", "Ausbildung",
                "Kenntnisse und Fähigkeiten", "Publikationen", "Zertifikate",
                "Sprachen", "Verfügbarkeit", "Links"
            ]
            
            if any(text.startswith(keyword) for keyword in section_keywords):
                paragraph.style = 'Heading 1'
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _format_lists(self):
        """Format lists to match LaTeX enumitem settings"""
        for paragraph in self.doc.paragraphs:
            if paragraph.text.strip().startswith('•') or paragraph.text.strip().startswith('-'):
                # This is a list item
                paragraph_format = paragraph.paragraph_format
                paragraph_format.left_indent = Pt(10)  # 10pt left margin
                paragraph_format.space_before = Pt(0)
                paragraph_format.space_after = Pt(-1)  # Negative spacing
                paragraph_format.line_spacing = 0.9
    
    def _handle_header_with_photo(self):
        """Handle the header table with photo"""
        # Look for existing header table created by Pandoc
        for i, table in enumerate(self.doc.tables):
            if table.rows and len(table.rows) >= 2 and len(table.rows[0].cells) >= 2:
                first_cell = table.rows[0].cells[0].text.strip()
                second_cell = table.rows[1].cells[0].text.strip()
                if "Arman Shirzad" in first_cell and "Cottbus" in second_cell:
                    self._fix_header_content(table)
                    self._add_photo_to_header_table(table)
                    return
    
    def _add_photo_to_header_table(self, table):
        """Add photo to existing header table"""
        # Try multiple photo paths
        photo_paths = [
            "presidency photo.png",  # Original location
            os.path.join(os.path.dirname(self.tex_file), "presidency photo.png"),  # Same directory as tex file
            os.path.join(os.getcwd(), "presidency photo.png"),  # Current working directory
        ]
        
        photo_found = False
        for photo_path in photo_paths:
            if os.path.exists(photo_path):
                try:
                    # Add photo to the second column of the first row
                    cell2 = table.rows[0].cells[1]
                    cell2_paragraph = cell2.paragraphs[0]
                    cell2_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    # Clear any existing content
                    cell2_paragraph.clear()
                    
                    # Add image
                    run = cell2_paragraph.add_run()
                    run.add_picture(photo_path, width=Cm(3), height=Cm(4))
                    
                    print(f"✓ Photo embedded successfully from: {photo_path}")
                    photo_found = True
                    break
                except Exception as e:
                    print(f"Warning: Could not embed photo from {photo_path}: {e}")
                    continue
        
        if not photo_found:
            print("Warning: Photo file 'presidency photo.png' not found in any expected location")
            print(f"Searched paths: {photo_paths}")
    
    def _format_hyperlinks(self):
        """Format hyperlinks to match LaTeX hyperref settings"""
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                if run.font.color.rgb and run.font.color.rgb != (0, 0, 0):
                    # Reset hyperlink color to black
                    run.font.color.rgb = None
    
    def _fix_header_content(self, table):
        """Fix header content to ensure phone number and proper formatting"""
        # Clean up the entire table structure
        if len(table.rows) > 1:
            # Remove extra rows - we only need the first row
            for i in range(len(table.rows) - 1, 0, -1):
                table._element.remove(table.rows[i]._element)
        
        # Fix the header cell content
        header_cell = table.rows[0].cells[0]
        header_text = header_cell.text
        
        # Clean up the header text and ensure proper formatting
        lines = header_text.split('\n')
        cleaned_lines = []
        
        # Process each line
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Fix LaTeX commands
            line = line.replace('\\quad', '    ')  # Replace \quad with spaces
            
            # If this is the contact line, ensure it has all components
            if "Cottbus" in line and "shirzarm@b-tu.de" in line:
                if "+49 157 5669 3804" not in line:
                    line = "Cottbus, Germany    shirzarm@b-tu.de    +49 157 5669 3804"
                else:
                    line = line.replace('\\quad', '    ')
            elif "Arman Shirzad" in line:
                # Keep the name line as is
                pass
            else:
                # Skip duplicate or malformed lines
                continue
                
            cleaned_lines.append(line)
        
        # Ensure we have the name and contact info
        if not any("Arman Shirzad" in line for line in cleaned_lines):
            cleaned_lines.insert(0, "Arman Shirzad")
        
        if not any("Cottbus" in line for line in cleaned_lines):
            cleaned_lines.append("Cottbus, Germany    shirzarm@b-tu.de    +49 157 5669 3804")
        
        header_cell.text = '\n'.join(cleaned_lines)
    
    def _fix_text_issues(self):
        """Fix various text formatting issues"""
        # Process all paragraphs
        for paragraph in self.doc.paragraphs:
            text = paragraph.text
            
            # Fix truncated dates (e.g., "/2022" -> "08/2022")
            text = self._fix_truncated_dates(text)
            
            # Fix percentage spacing (e.g., "∼30%" -> "∼30 %")
            text = self._fix_percentage_spacing(text)
            
            # Remove stray page numbers
            text = self._remove_stray_page_numbers(text)
            
            if text != paragraph.text:
                paragraph.text = text
        
        # Also fix text in tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        
                        # Fix truncated dates
                        text = self._fix_truncated_dates(text)
                        
                        # Fix percentage spacing
                        text = self._fix_percentage_spacing(text)
                        
                        # Remove stray page numbers
                        text = self._remove_stray_page_numbers(text)
                        
                        if text != paragraph.text:
                            paragraph.text = text
        
        # Additional comprehensive text replacement
        self._comprehensive_text_replacement()
        
        # Final run-level text replacement
        self._fix_runs_directly()
        
        # Fix missing content and symbols
        self._fix_missing_content()
        
        # Clean up any remaining duplicates
        self._clean_duplicates()
    
    def _fix_truncated_dates(self, text):
        """Fix truncated month digits in dates"""
        import re
        
        # More comprehensive date fixing
        # Pattern 1: "/2022 bis 03/2025" -> "08/2022 bis 03/2025" (Refah Bank)
        text = re.sub(r'/(2022) bis (03)/(2025)', r'08/\1 bis \2/\3', text)
        
        # Pattern 2: "/2021 bis 07/2022" -> "03/2021 bis 07/2022" (MAPSA)
        text = re.sub(r'/(2021) bis (07)/(2022)', r'03/\1 bis \2/\3', text)
        
        # Pattern 3: "/2020 bis heute" -> "08/2020 bis heute" (Freelance)
        text = re.sub(r'/(2020) bis heute', r'08/\1 bis heute', text)
        
        # Pattern 4: "/2025 bis heute" -> "03/2025 bis heute" (Master)
        text = re.sub(r'/(2025) bis heute', r'03/\1 bis heute', text)
        
        # Pattern 5: "/2016 bis 10/2020" -> "09/2016 bis 10/2020" (Bachelor)
        text = re.sub(r'/(2016) bis (10)/(2020)', r'09/\1 bis \2/\3', text)
        
        return text
    
    def _fix_percentage_spacing(self, text):
        """Fix spacing before percentage symbols"""
        import re
        # Fix "∼30%" to "∼30 %" (add space before %)
        text = re.sub(r'∼(\d+)%', r'∼\1 %', text)
        return text
    
    def _remove_stray_page_numbers(self, text):
        """Remove stray page numbers that appear in content"""
        import re
        # Remove standalone numbers that might be page numbers
        # This is a conservative approach - only remove obvious stray numbers
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Skip lines that are just a single digit (likely stray page numbers)
            if re.match(r'^\s*\d\s*$', line.strip()):
                continue
            cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _comprehensive_text_replacement(self):
        """Comprehensive text replacement to fix all remaining issues"""
        import re
        
        # Define all the specific replacements needed
        replacements = [
            # Fix duplicated dates first - more specific patterns
            (r'0808/(\d{4})', r'08/\1'),  # "0808/2022" -> "08/2022"
            (r'0303/(\d{4})', r'03/\1'),  # "0303/2021" -> "03/2021"
            (r'0708/(\d{4})', r'07/\1'),  # "0708/2022" -> "07/2022"
            (r'0909/(\d{4})', r'09/\1'),  # "0909/2016" -> "09/2016"
            (r'1008/(\d{4})', r'10/\1'),  # "1008/2020" -> "10/2020"
            
            # Then fix any remaining truncated dates
            (r'/(2022)', '08/2022'),
            (r'/(2021)', '03/2021'), 
            (r'/(2020)', '08/2020'),
            (r'/(2025)', '03/2025'),
            (r'/(2016)', '09/2016'),
        ]
        
        # Apply replacements to all paragraphs
        for paragraph in self.doc.paragraphs:
            text = paragraph.text
            original_text = text
            
            for pattern, replacement in replacements:
                text = re.sub(pattern, replacement, text)
            
            if text != original_text:
                paragraph.text = text
        
        # Apply replacements to all table cells
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        original_text = text
                        
                        for pattern, replacement in replacements:
                            text = re.sub(pattern, replacement, text)
                        
                        if text != original_text:
                            paragraph.text = text
    
    def _fix_runs_directly(self):
        """Fix text issues directly in runs"""
        import re
        
        # Define replacements for runs
        replacements = [
            (r'0808/(\d{4})', r'08/\1'),  # "0808/2022" -> "08/2022"
            (r'0303/(\d{4})', r'03/\1'),  # "0303/2021" -> "03/2021"
            (r'0708/(\d{4})', r'07/\1'),  # "0708/2022" -> "07/2022"
            (r'0909/(\d{4})', r'09/\1'),  # "0909/2016" -> "09/2016"
            (r'1008/(\d{4})', r'10/\1'),  # "1008/2020" -> "10/2020"
        ]
        
        # Fix runs in paragraphs
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                text = run.text
                original_text = text
                
                for pattern, replacement in replacements:
                    text = re.sub(pattern, replacement, text)
                
                if text != original_text:
                    run.text = text
        
        # Fix runs in table cells
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            text = run.text
                            original_text = text
                            
                            for pattern, replacement in replacements:
                                text = re.sub(pattern, replacement, text)
                            
                            if text != original_text:
                                run.text = text
    
    def _fix_missing_content(self):
        """Fix missing content and symbols throughout the document"""
        import re
        
        # Define content fixes
        content_fixes = [
            # Fix missing symbols and content - be more specific to avoid duplicates
            (r' ∼ 10,000 ∼ 10,000 Requests/Tag', ' ∼ 10,000 Requests/Tag'),
            (r' ∼ 200,000 ∼ 200,000 Nutzer', ' ∼ 200,000 Nutzer'),
            (r'Kosten ∼ 30 % ∼ 30 % niedriger', 'Kosten ∼ 30 % niedriger'),
            (r' Requests/Tag', ' ∼ 10,000 Requests/Tag'),
            (r' Nutzer integriert', ' ∼ 200,000 Nutzer integriert'),
            (r'Kosten  niedriger', 'Kosten ∼ 30 % niedriger'),
            (r'\\quad', '    '),  # Replace LaTeX quad with spaces
            (r'\\sim', '≈'),       # Replace LaTeX sim with Unicode
            (r'\\approx', '≈'),    # Replace LaTeX approx with Unicode
        ]
        
        # Apply fixes to all paragraphs
        for paragraph in self.doc.paragraphs:
            text = paragraph.text
            original_text = text
            
            for pattern, replacement in content_fixes:
                text = re.sub(pattern, replacement, text)
            
            if text != original_text:
                paragraph.text = text
        
        # Apply fixes to all table cells
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        original_text = text
                        
                        for pattern, replacement in content_fixes:
                            text = re.sub(pattern, replacement, text)
                        
                        if text != original_text:
                            paragraph.text = text
        
        # Apply fixes to runs as well
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                text = run.text
                original_text = text
                
                for pattern, replacement in content_fixes:
                    text = re.sub(pattern, replacement, text)
                
                if text != original_text:
                    run.text = text
        
        # Apply fixes to runs in tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            text = run.text
                            original_text = text
                            
                            for pattern, replacement in content_fixes:
                                text = re.sub(pattern, replacement, text)
                            
                            if text != original_text:
                                run.text = text
    
    def _clean_duplicates(self):
        """Clean up duplicate content and formatting issues"""
        import re
        
        # Define duplicate patterns to fix
        duplicate_patterns = [
            (r'(∼ \d+,?\d*)\s+\1', r'\1'),  # Remove duplicate numbers like "∼ 10,000 ∼ 10,000"
            (r'(Cottbus, Germany)\s+\1', r'\1'),  # Remove duplicate location
            (r'(shirzarm@b-tu\.de)\s+\1', r'\1'),  # Remove duplicate email
            (r'(Arman Shirzad)\s+\1', r'\1'),  # Remove duplicate name
        ]
        
        # Apply to all paragraphs
        for paragraph in self.doc.paragraphs:
            text = paragraph.text
            original_text = text
            
            for pattern, replacement in duplicate_patterns:
                text = re.sub(pattern, replacement, text)
            
            if text != original_text:
                paragraph.text = text
        
        # Apply to all table cells
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        original_text = text
                        
                        for pattern, replacement in duplicate_patterns:
                            text = re.sub(pattern, replacement, text)
                        
                        if text != original_text:
                            paragraph.text = text
    
    def save_document(self):
        """Save the final document"""
        print(f"Saving document as {self.docx_file}...")
        self.doc.save(self.docx_file)
        print(f"✓ Document saved as {self.docx_file}")
        
        # Clean up temporary file
        temp_docx_path = os.path.join(os.path.dirname(self.docx_file), "temp_" + os.path.basename(self.docx_file))
        if os.path.exists(temp_docx_path):
            os.remove(temp_docx_path)
    
    def convert(self):
        """Main conversion process"""
        print("Starting LaTeX to Word conversion...")
        
        if not os.path.exists(self.tex_file):
            print(f"Error: LaTeX file '{self.tex_file}' not found")
            return False
        
        # Step 1: Convert with Pandoc
        if not self.convert_with_pandoc():
            return False
        
        # Step 2: Load and refine
        self.load_and_refine_document()
        
        # Step 3: Save final document
        self.save_document()
        
        print("✓ Conversion completed successfully!")
        print(f"Output file: {self.docx_file}")
        return True

def main():
    """Main function"""
    import sys
    
    # Check if specific files are provided as arguments
    if len(sys.argv) > 1:
        input_files = sys.argv[1:]
    else:
        # Default files to convert
        input_files = ["cvde.tex", "cven.tex"]
    
    success_count = 0
    total_files = len(input_files)
    
    for tex_file in input_files:
        if not os.path.exists(tex_file):
            print(f"❌ Error: LaTeX file '{tex_file}' not found")
            continue
            
        print(f"\n{'='*60}")
        print(f"Converting: {tex_file}")
        print(f"{'='*60}")
        
        # Generate output filename
        base_name = os.path.splitext(tex_file)[0]
        docx_file = f"{base_name}.docx"
        
        converter = LaTeXToWordConverter(tex_file, docx_file)
        success = converter.convert()
        
        if success:
            success_count += 1
            print(f"\n✅ SUCCESS: {tex_file} → {docx_file}")
        else:
            print(f"\n❌ FAILED: {tex_file}")
    
    print(f"\n{'='*60}")
    print(f"CONVERSION SUMMARY")
    print(f"{'='*60}")
    print(f"Successfully converted: {success_count}/{total_files} files")
    
    if success_count > 0:
        print("\nNext steps:")
        print("1. Open the Word documents")
        print("2. Review formatting and make minor adjustments if needed")
        print("3. Save and export to PDF to compare with original LaTeX PDF")
        print("\nNote: Some font differences may occur (Computer Modern → Liberation Serif)")
        print("This is normal and expected for cross-platform compatibility.")
    
    if success_count == 0:
        print("All conversions failed. Please check the error messages above.")
        sys.exit(1)

if __name__ == "__main__":
    main()
